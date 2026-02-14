import os
import threading
import zipfile
import logging
from io import BytesIO
from functools import wraps

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_file,
    flash
)

from flask_login import LoginManager, login_required, current_user
from werkzeug.utils import secure_filename
from docx import Document as WordDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

from config import Config
from models import db, User, Document, ChatHistory
from auth import auth
from ocr_pipeline import run_ocr
from rag_engine import ask_question


# ======================================================
# APP SETUP
# ======================================================

app = Flask(__name__)
app.config.from_object(Config)

db.init_app(app)

login_manager = LoginManager()
login_manager.login_view = "auth.login"
login_manager.init_app(app)

app.register_blueprint(auth)


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ======================================================
# LOGGING
# ======================================================

logging.basicConfig(
    filename="system.log",
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)


# ======================================================
# ROLE DECORATOR
# ======================================================

def role_required(role):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            if current_user.role != role:
                flash("Access denied.")
                return redirect(url_for("dashboard"))
            return func(*args, **kwargs)
        return wrapper
    return decorator


# ======================================================
# BACKGROUND OCR
# ======================================================

def process_ocr_background(app, doc_id, path):
    with app.app_context():
        doc = Document.query.get(doc_id)

        if not doc:
            return

        try:
            text = run_ocr(path)
            doc.extracted_text = text
            doc.status = "completed"
            logging.info(f"OCR completed for doc {doc_id}")
        except Exception as e:
            doc.status = "failed"
            doc.extracted_text = f"OCR Error: {str(e)}"
            logging.error(f"OCR failed for doc {doc_id}: {e}")

        db.session.commit()


# ======================================================
# DASHBOARD (AUTO ROLE SWITCH)
# ======================================================

@app.route("/")
@login_required
def dashboard():
    print("Current user:", current_user.email, current_user.role)
    # ADMIN DASHBOARD
    if current_user.role == "admin":

        users = User.query.order_by(User.id.desc()).all()
        documents = Document.query.order_by(Document.id.desc()).all()

        total_users = User.query.count()
        total_docs = Document.query.count()
        total_completed = Document.query.filter_by(status="completed").count()
        total_failed = Document.query.filter_by(status="failed").count()

        return render_template(
            "admin_dashboard.html",
            users=users,
            documents=documents,
            total_users=total_users,
            total_docs=total_docs,
            total_completed=total_completed,
            total_failed=total_failed
        )

    # NORMAL USER DASHBOARD
    documents = Document.query.filter_by(
        user_id=current_user.id
    ).order_by(Document.id.desc()).all()

    return render_template(
        "dashboard.html",
        user=current_user,
        documents=documents
    )


# ======================================================
# UPLOAD
# ======================================================

@app.route("/upload", methods=["POST"])
@login_required
def upload():

    files = request.files.getlist("documents")

    if not files:
        flash("No files selected.")
        return redirect(url_for("dashboard"))

    for file in files:
        if file and file.filename:

            filename = secure_filename(file.filename)
            save_path = os.path.join(
                app.config["UPLOAD_FOLDER"],
                filename
            )

            file.save(save_path)

            doc = Document(
                filename=filename,
                stored_path=save_path,
                user_id=current_user.id,
                status="processing"
            )

            db.session.add(doc)
            db.session.commit()

            thread = threading.Thread(
                target=process_ocr_background,
                args=(app, doc.id, save_path)
            )
            thread.daemon = True
            thread.start()

    flash("Upload successful. OCR is processing.")
    return redirect(url_for("dashboard"))


# ======================================================
# VIEW DOCUMENT
# ======================================================

@app.route("/document/<int:doc_id>", methods=["GET", "POST"])
@login_required
def view_document(doc_id):

    doc = Document.query.get_or_404(doc_id)

    if doc.user_id != current_user.id and current_user.role != "admin":
        return "Unauthorized", 403

    chat_history = ChatHistory.query.filter_by(
        document_id=doc.id
    ).order_by(ChatHistory.id.desc()).all()

    if request.method == "POST":
        edited_text = request.form.get("edited_text")
        doc.extracted_text = edited_text
        db.session.commit()
        flash("Document updated.")
        return redirect(url_for("view_document", doc_id=doc.id))

    return render_template(
        "view_document.html",
        doc=doc,
        chat_history=chat_history
    )


# ======================================================
# ASK QUESTION (RAG)
# ======================================================
@app.route("/ask/<int:doc_id>", methods=["POST"])
@login_required
def ask_route(doc_id):

    doc = Document.query.get_or_404(doc_id)

    # Allow admin to access any document
    if doc.user_id != current_user.id and current_user.role != "admin":
        return "Unauthorized", 403

    question = request.form.get("question", "").strip()

    if not question:
        flash("Please enter a question.", "warning")
        return redirect(url_for("view_document", doc_id=doc.id))

    if not doc.extracted_text:
        flash("Document has no extracted text.", "danger")
        return redirect(url_for("view_document", doc_id=doc.id))

    try:
        answer, citation = ask_question(doc.extracted_text, question)

        chat = ChatHistory(
            document_id=doc.id,
            question=question,
            answer=answer,
            citation=citation
        )

        db.session.add(chat)
        db.session.commit()

    except Exception as e:
        flash("Error while processing question.", "danger")
        print("RAG ERROR:", e)

    return redirect(url_for("view_document", doc_id=doc.id))

# ======================================================
# DOWNLOADS
# ======================================================

@app.route("/download/pdf/<int:doc_id>")
@login_required
def download_pdf(doc_id):

    doc = Document.query.get_or_404(doc_id)

    if doc.user_id != current_user.id:
        return "Unauthorized", 403

    os.makedirs("outputs", exist_ok=True)
    output_path = os.path.join("outputs", f"{doc.id}.pdf")

    pdf = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    for line in (doc.extracted_text or "").split("\n"):
        elements.append(Paragraph(line, styles["Normal"]))

    pdf.build(elements)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{doc.filename}.pdf"
    )


@app.route("/download/word/<int:doc_id>")
@login_required
def download_word(doc_id):

    doc = Document.query.get_or_404(doc_id)

    if doc.user_id != current_user.id:
        return "Unauthorized", 403

    os.makedirs("outputs", exist_ok=True)
    output_path = os.path.join("outputs", f"{doc.id}.docx")

    word_doc = WordDocument()
    word_doc.add_paragraph(doc.extracted_text or "")
    word_doc.save(output_path)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{doc.filename}.docx"
    )


# ======================================================
# BULK ZIP DOWNLOAD
# ======================================================

@app.route("/download/bulk", methods=["POST"])
@login_required
def bulk_download():

    selected_ids = request.form.getlist("selected_docs")

    if not selected_ids:
        flash("No documents selected.")
        return redirect(url_for("dashboard"))

    memory_file = BytesIO()

    with zipfile.ZipFile(memory_file, "w") as zf:
        for doc_id in selected_ids:
            doc = Document.query.get(int(doc_id))
            if doc and doc.user_id == current_user.id:
                zf.writestr(
                    f"{doc.filename}.txt",
                    doc.extracted_text or ""
                )

    memory_file.seek(0)

    return send_file(
        memory_file,
        download_name="documents.zip",
        as_attachment=True
    )


# ======================================================
# USER DELETE DOCUMENT
# ======================================================

@app.route("/delete/<int:doc_id>")
@login_required
def delete_document_user(doc_id):

    doc = Document.query.get_or_404(doc_id)

    if doc.user_id != current_user.id:
        return "Unauthorized", 403

    db.session.delete(doc)
    db.session.commit()

    flash("Document deleted.")
    return redirect(url_for("dashboard"))


# ======================================================
# ADMIN DELETE USER
# ======================================================

@app.route("/admin/delete_user/<int:user_id>")
@login_required
@role_required("admin")
def delete_user(user_id):

    user = User.query.get_or_404(user_id)

    if user.role == "admin":
        flash("Cannot delete admin.")
        return redirect(url_for("dashboard"))

    Document.query.filter_by(user_id=user.id).delete()
    db.session.delete(user)
    db.session.commit()

    flash("User deleted.")
    return redirect(url_for("dashboard"))


# ======================================================
# ADMIN DELETE DOCUMENT
# ======================================================

@app.route("/admin/delete_document/<int:doc_id>")
@login_required
@role_required("admin")
def delete_document_admin(doc_id):

    doc = Document.query.get_or_404(doc_id)

    # 🔥 DELETE CHAT HISTORY FIRST
    ChatHistory.query.filter_by(document_id=doc.id).delete()

    # 🔥 Then delete document
    db.session.delete(doc)
    db.session.commit()

    flash("Document deleted successfully.")
    return redirect(url_for("dashboard"))

# ======================================================
# RUN
# ======================================================

if __name__ == "__main__":
    with app.app_context():
        db.create_all()

    app.run(host="0.0.0.0", port=5000, debug=True)